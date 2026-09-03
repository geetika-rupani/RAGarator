"""Load PDF, DOCX, and TXT files into raw text."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.utils.errors import ExtractionError, UnsupportedFileError


def load_document(path: Path) -> str:
    """Extract raw text from a supported document path."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".txt":
            return _load_txt(path)
        if suffix == ".pdf":
            return _load_pdf(path)
        if suffix == ".docx":
            return _load_docx(path)
    except UnsupportedFileError:
        raise
    except Exception as exc:
        raise ExtractionError(f"Failed to extract text from {path.name}: {exc}") from exc
    raise UnsupportedFileError(
        f"Unsupported file type '{suffix}'. Upload a PDF, DOCX, or TXT file."
    )


def _load_txt(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        pages.append(extracted if extracted.strip() else "")
    if not any(page.strip() for page in pages):
        raise ExtractionError(
            f"{path.name} appears to be an image-only PDF with no extractable text."
        )
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    document = Document(str(path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    if not blocks:
        raise ExtractionError(f"{path.name} contains no readable paragraphs or tables.")
    return "\n".join(blocks)
